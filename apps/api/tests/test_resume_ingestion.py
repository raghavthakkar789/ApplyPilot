import hashlib
import io
import os
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import pymupdf
import pytest
from docx import Document
from sqlalchemy import select, text
from sqlalchemy.exc import DatabaseError

from applypilot.core.config import get_settings
from applypilot.models.candidate_fact import CandidateFactVersion
from applypilot.models.resume import ResumeVersion, StoredDocument
from applypilot.parsers.base import DocumentParseError
from applypilot.repositories.database import SessionFactory
from applypilot.services.document_extraction_service import parser_deadline
from applypilot.services.document_storage_service import DocumentStorageService
from tests.auth_helpers import ORIGIN, client, initialize, reset_auth_database


@pytest.fixture(autouse=True)
def clean_database() -> None:
    reset_auth_database()


def csrf_headers(owner: object) -> dict[str, str]:
    token = owner.cookies.get("applypilot_csrf")  # type: ignore[attr-defined]
    assert token
    return {**ORIGIN, "X-ApplyPilot-CSRF": token}


def pdf_bytes(text_value: str = "Skill: Python") -> bytes:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text_value)
    value = document.tobytes()
    document.close()
    return value


def docx_bytes(text_value: str = "Language: English") -> bytes:
    stream = io.BytesIO()
    document = Document()
    document.add_paragraph(text_value)
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Structured table text"
    document.save(stream)
    return stream.getvalue()


def upload(owner: object, name: str, content: bytes, media_type: str):  # type: ignore[no-untyped-def]
    return owner.post(  # type: ignore[attr-defined,no-any-return]
        "/api/resumes",
        headers=csrf_headers(owner),
        data={"display_name": "Synthetic resume", "purpose": "Test fixture"},
        files={"file": (name, content, media_type)},
    )


@pytest.mark.parametrize(
    ("name", "content", "media_type", "expected"),
    [
        ("sample.pdf", pdf_bytes(), "application/pdf", "pdf"),
        (
            "sample.docx",
            docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        ),
        ("sample.txt", b"Skill: Python\n", "text/plain", "text"),
    ],
)
def test_supported_uploads_extract_and_hash(
    name: str, content: bytes, media_type: str, expected: str
) -> None:
    with client() as owner:
        initialize(owner)
        response = upload(owner, name, content, media_type)
        assert response.status_code == 201, response.text
        version = response.json()["current_version"]
        assert version["format"] == expected
        assert version["sha256"] == hashlib.sha256(content).hexdigest()
        extraction = owner.get(f"/api/resumes/versions/{version['id']}/extraction")
        assert extraction.status_code == 200
        assert extraction.json()["status"] == "succeeded"


def test_authentication_csrf_and_origin_are_required() -> None:
    with client() as anonymous:
        assert anonymous.get("/api/resumes").status_code == 401
    with client() as owner:
        initialize(owner)
        files = {"file": ("sample.txt", b"safe", "text/plain")}
        data = {"display_name": "Synthetic"}
        assert owner.post("/api/resumes", data=data, files=files).status_code == 403
        token = owner.cookies.get("applypilot_csrf")
        assert token
        wrong = {"Origin": "http://localhost:3000", "X-ApplyPilot-CSRF": token}
        assert owner.post("/api/resumes", headers=wrong, data=data, files=files).status_code == 403


@pytest.mark.parametrize(
    ("name", "content", "media_type"),
    [
        ("resume.pdf", b"not a pdf", "application/pdf"),
        ("resume.txt", b"%PDF-fake", "text/plain"),
        ("resume.exe", b"MZ", "application/octet-stream"),
        ("resume.txt", b"text\x00binary", "text/plain"),
        ("resume.txt", b"\xff\xfeinvalid", "text/plain"),
        ("resume.docx", b"PK\x03\x04invalid", "application/zip"),
    ],
)
def test_unsafe_or_ambiguous_uploads_are_rejected(
    name: str, content: bytes, media_type: str
) -> None:
    with client() as owner:
        initialize(owner)
        assert upload(owner, name, content, media_type).status_code == 422


def test_oversized_upload_is_rejected() -> None:
    with client() as owner:
        initialize(owner)
        response = upload(owner, "large.txt", b"a" * (10 * 1024 * 1024 + 1), "text/plain")
        assert response.status_code == 413
        temporary = Path(get_settings().document_storage_root) / "temporary"
        assert not list(temporary.glob("*.upload"))


def test_encrypted_and_empty_pdf_behavior() -> None:
    encrypted = pymupdf.open()
    encrypted.new_page()
    value = encrypted.tobytes(
        encryption=pymupdf.PDF_ENCRYPT_AES_256,
        owner_pw="synthetic-owner",
        user_pw="synthetic-user",
    )
    encrypted.close()
    with client() as owner:
        initialize(owner)
        assert upload(owner, "locked.pdf", value, "application/pdf").status_code == 422
        empty = upload(owner, "empty.pdf", pdf_bytes(""), "application/pdf")
        extraction = owner.get(
            f"/api/resumes/versions/{empty.json()['current_version']['id']}/extraction"
        ).json()
        assert any("OCR was not run" in warning for warning in extraction["warnings"])


def test_docx_traversal_and_zip_bomb_controls() -> None:
    traversal = io.BytesIO()
    with zipfile.ZipFile(traversal, "w") as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "document")
        archive.writestr("../escape", "unsafe")
    compressed = io.BytesIO()
    with zipfile.ZipFile(compressed, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "types")
        archive.writestr("word/document.xml", "0" * 1_000_000)
    media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    with client() as owner:
        initialize(owner)
        assert upload(owner, "bad.docx", traversal.getvalue(), media).status_code == 422
        assert upload(owner, "bomb.docx", compressed.getvalue(), media).status_code == 422


def test_pdf_page_and_extracted_character_limits() -> None:
    document = pymupdf.open()
    for _ in range(251):
        document.new_page()
    too_many_pages = document.tobytes()
    document.close()
    with client() as owner:
        initialize(owner)
        assert upload(owner, "pages.pdf", too_many_pages, "application/pdf").status_code == 422
        too_much_text = b"a" * 2_000_001
        assert upload(owner, "long.txt", too_much_text, "text/plain").status_code == 422


def test_docx_external_relationship_is_not_followed() -> None:
    source = io.BytesIO(docx_bytes())
    output = io.BytesIO()
    with zipfile.ZipFile(source) as original, zipfile.ZipFile(output, "w") as changed:
        for item in original.infolist():
            data = original.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                data = data.replace(
                    b"</Relationships>",
                    b'<Relationship Id="external" Type="link" Target="https://invalid.example/" '
                    b'TargetMode="External"/></Relationships>',
                )
            changed.writestr(item, data)
    media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    with client() as owner:
        initialize(owner)
        response = upload(owner, "external.docx", output.getvalue(), media)
        assert response.status_code == 201
        version_id = response.json()["current_version"]["id"]
        warnings = owner.get(f"/api/resumes/versions/{version_id}/extraction").json()["warnings"]
        assert warnings == ["External document relationships were ignored."]


def test_storage_replaces_symlink_without_following_it(tmp_path: Path) -> None:
    storage = DocumentStorageService(str(tmp_path / "storage"))
    external = tmp_path / "outside"
    external.write_text("unchanged")
    target = tmp_path / "storage" / "originals" / "safe.txt"
    target.symlink_to(external)
    staged = tmp_path / "storage" / "temporary" / "staged"
    staged.write_text("stored")
    storage.finalize(staged, "safe.txt")
    assert not target.is_symlink()
    assert target.read_text() == "stored"
    assert external.read_text() == "unchanged"
    assert oct(os.stat(target).st_mode & 0o777) == "0o600"


def test_parser_deadline_fails_closed() -> None:
    with pytest.raises(DocumentParseError, match="safety timeout"):
        with parser_deadline(0.01):
            time.sleep(0.1)


def test_duplicate_storage_version_history_and_download_headers() -> None:
    content = b"Skill: Python\n"
    with client() as owner:
        initialize(owner)
        first = upload(owner, "../resume.txt", content, "text/plain")
        assert first.status_code == 201
        resume_id = first.json()["id"]
        second = owner.post(
            f"/api/resumes/{resume_id}/versions",
            headers=csrf_headers(owner),
            files={"file": ("next.txt", content, "text/plain")},
        )
        assert second.status_code == 200
        detail = owner.get(f"/api/resumes/{resume_id}").json()
        assert [item["version_number"] for item in detail["versions"]] == [2, 1]
        with SessionFactory() as database:
            assert len(list(database.scalars(select(StoredDocument)))) == 1
            original = database.scalar(
                select(ResumeVersion).where(ResumeVersion.version_number == 1)
            )
            assert original
            with pytest.raises(DatabaseError):
                database.execute(
                    text("UPDATE resume_versions SET document_id = 'changed' WHERE id = :id"),
                    {"id": original.id},
                )
                database.flush()
            database.rollback()
        version_id = detail["versions"][0]["id"]
        download = owner.get(f"/api/resumes/versions/{version_id}/download")
        assert download.status_code == 200
        assert download.headers["x-content-type-options"] == "nosniff"
        assert "no-store" in download.headers["cache-control"]
        assert "attachment" in download.headers["content-disposition"]
    with client() as anonymous:
        assert anonymous.get(f"/api/resumes/versions/{version_id}/download").status_code == 401


def test_candidate_acceptance_remains_unverified_and_rejection_creates_no_fact() -> None:
    with client() as owner:
        initialize(owner)
        created = upload(owner, "resume.txt", b"Skill: Python\nLanguage: English\n", "text/plain")
        version_id = created.json()["current_version"]["id"]
        candidates = owner.get(f"/api/resume-versions/{version_id}/fact-candidates").json()[
            "candidates"
        ]
        accepted = owner.post(
            f"/api/resume-fact-candidates/{candidates[0]['id']}/accept",
            headers=csrf_headers(owner),
        )
        assert accepted.json()["fact_state"] == "unverified"
        fact = owner.get(f"/api/candidate-facts/{accepted.json()['fact_identity_id']}").json()
        assert fact["current_version"]["lifecycle_state"] == "unverified"
        owner.post(
            f"/api/resume-fact-candidates/{candidates[1]['id']}/reject",
            headers=csrf_headers(owner),
        )
        with SessionFactory() as database:
            assert len(list(database.scalars(select(CandidateFactVersion)))) == 1


def test_trash_restore_and_dependency_aware_deletion() -> None:
    with client() as owner:
        initialize(owner)
        created = upload(owner, "resume.txt", b"Skill: Python\n", "text/plain").json()
        resume_id = created["id"]
        version_id = created["current_version"]["id"]
        assert (
            owner.post(f"/api/resumes/{resume_id}/trash", headers=csrf_headers(owner)).status_code
            == 200
        )
        assert (
            owner.post(f"/api/resumes/{resume_id}/restore", headers=csrf_headers(owner)).status_code
            == 200
        )
        owner.post(f"/api/resumes/{resume_id}/trash", headers=csrf_headers(owner))
        candidate_id = owner.get(f"/api/resume-versions/{version_id}/fact-candidates").json()[
            "candidates"
        ][0]["id"]
        owner.post(
            f"/api/resume-fact-candidates/{candidate_id}/accept", headers=csrf_headers(owner)
        )
        assert (
            owner.delete(f"/api/resumes/{resume_id}", headers=csrf_headers(owner)).status_code
            == 409
        )


def test_permanent_deletion_removes_unreferenced_private_content() -> None:
    with client() as owner:
        initialize(owner)
        created = upload(owner, "resume.txt", b"Synthetic narrative only\n", "text/plain").json()
        resume_id = created["id"]
        version_id = created["current_version"]["id"]
        with SessionFactory() as database:
            version = database.get(ResumeVersion, version_id)
            assert version
            document = database.get(StoredDocument, version.document_id)
            assert document
            path = DocumentStorageService(get_settings().document_storage_root).original_path(
                document.storage_key
            )
            assert path.exists()
        owner.post(f"/api/resumes/{resume_id}/trash", headers=csrf_headers(owner))
        assert (
            owner.delete(f"/api/resumes/{resume_id}", headers=csrf_headers(owner)).status_code
            == 204
        )
        assert not path.exists()
        assert owner.get(f"/api/resumes/versions/{version_id}/download").status_code == 404


def test_concurrent_version_allocation_keeps_unique_numbers() -> None:
    with client() as owner:
        initialize(owner)
        resume_id = upload(owner, "resume.txt", b"base", "text/plain").json()["id"]
        cookies = dict(owner.cookies)

    def attempt(index: int) -> int:
        with client() as concurrent:
            concurrent.cookies.update(cookies)
            response = concurrent.post(
                f"/api/resumes/{resume_id}/versions",
                headers=csrf_headers(concurrent),
                files={"file": (f"v{index}.txt", f"value {index}".encode(), "text/plain")},
            )
            return response.status_code

    with ThreadPoolExecutor(max_workers=20) as pool:
        statuses = list(pool.map(attempt, range(20)))
    assert all(status in {200, 409} for status in statuses)
    with SessionFactory() as database:
        versions = list(
            database.scalars(select(ResumeVersion).where(ResumeVersion.resume_id == resume_id))
        )
        numbers = [version.version_number for version in versions]
        assert len(numbers) == len(set(numbers))
