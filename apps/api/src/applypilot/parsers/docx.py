from pathlib import Path, PurePosixPath
from zipfile import BadZipFile, ZipFile

from docx import Document

from applypilot.domain.documents.document_limits import LIMITS
from applypilot.domain.documents.extraction_result import EvidenceSegment, ExtractionResult
from applypilot.parsers.base import DocumentParseError


class DocxParser:
    version = "python-docx-1.2.0"

    def parse(self, path: Path) -> ExtractionResult:
        warnings: list[str] = []
        self._validate_archive(path, warnings)
        try:
            document = Document(str(path))
        except Exception as error:
            raise DocumentParseError("The DOCX structure is invalid.") from error
        segments: list[EvidenceSegment] = []
        for index, paragraph in enumerate(document.paragraphs, 1):
            if paragraph.text.strip():
                segments.append(EvidenceSegment(f"paragraph {index}", paragraph.text))
        for table_index, table in enumerate(document.tables, 1):
            for row_index, row in enumerate(table.rows, 1):
                value = " | ".join(cell.text for cell in row.cells).strip()
                if value:
                    segments.append(EvidenceSegment(f"table {table_index}, row {row_index}", value))
        text = "\n".join(segment.text for segment in segments)
        if len(text) > LIMITS.extracted_characters:
            raise DocumentParseError("The extracted text exceeds the safety limit.")
        return ExtractionResult(
            text=text,
            segments=tuple(segments),
            paragraph_count=len(segments),
            warnings=tuple(warnings),
        )

    def _validate_archive(self, path: Path, warnings: list[str]) -> None:
        try:
            with ZipFile(path) as archive:
                entries = archive.infolist()
                if len(entries) > LIMITS.docx_entries:
                    raise DocumentParseError("The DOCX has too many archive entries.")
                total = 0
                names = {entry.filename for entry in entries}
                for entry in entries:
                    pure = PurePosixPath(entry.filename)
                    if pure.is_absolute() or ".." in pure.parts or "\\" in entry.filename:
                        raise DocumentParseError("The DOCX contains an unsafe archive path.")
                    total += entry.file_size
                    if total > LIMITS.docx_uncompressed_bytes:
                        raise DocumentParseError("The DOCX expands beyond the safety limit.")
                    if entry.file_size and entry.compress_size == 0:
                        raise DocumentParseError("The DOCX has an unsafe compression ratio.")
                    if (
                        entry.compress_size
                        and entry.file_size / entry.compress_size > LIMITS.docx_compression_ratio
                    ):
                        raise DocumentParseError("The DOCX has an unsafe compression ratio.")
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise DocumentParseError("The file is not a valid DOCX document.")
                content_types = archive.read("[Content_Types].xml")
                if b"macroEnabled" in content_types or any("vbaProject" in name for name in names):
                    raise DocumentParseError("Macro-enabled documents are not supported.")
                for name in names:
                    if name.endswith(".rels") and b'TargetMode="External"' in archive.read(name):
                        warnings.append("External document relationships were ignored.")
        except BadZipFile as error:
            raise DocumentParseError("The DOCX container is invalid.") from error
